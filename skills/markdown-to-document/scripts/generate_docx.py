"""
Generate a Word document from JSON or Markdown input.

Usage:
    python generate_docx.py input.json -t template.docx -o output.docx
    python generate_docx.py input.json -o output.docx
    python generate_docx.py input.md -o output.docx

Input formats:
    JSON with template: {"data": {"field": "value", "items": [...]}}
        Data is rendered into the Jinja2 template via docxtpl.
    JSON without template: {"data": {"title": "...", "sections": [{"heading": "...", "body": "..."}]}}
        Creates a document from scratch using python-docx.
    Markdown: Content is converted to a structured Word document.

Requires: pip install docxtpl (includes python-docx)
"""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path


def _ensure_package(import_name: str, pip_name: str | None = None) -> bool:
    """Try to import a package; auto-install if non-interactive. Returns True if available."""
    pip_name = pip_name or import_name
    try:
        __import__(import_name)
        return True
    except (ImportError, OSError):
        pass

    if sys.stdin.isatty():
        answer = input(f'Package "{pip_name}" is not installed. Install it now? [Y/n] ').strip().lower()
        if answer in ("", "y", "yes"):
            subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])
            return True
        return False
    else:
        print(f"Auto-installing missing package: {pip_name}", file=sys.stderr)
        subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])
        __import__(import_name)  # verify
        return True


def load_json(input_path: Path) -> dict:
    """Load and validate JSON input."""
    text = input_path.read_text(encoding="utf-8")
    data = json.loads(text)
    if "data" not in data:
        print("Error: JSON must have a 'data' key.", file=sys.stderr)
        sys.exit(1)
    return data["data"]


def generate_with_template(data: dict, template_path: Path, output_path: Path) -> None:
    """Render a docxtpl template with JSON data."""
    if not _ensure_package("docxtpl"):
        print("Error: docxtpl is required.", file=sys.stderr)
        sys.exit(1)
    from docxtpl import DocxTemplate

    doc = DocxTemplate(str(template_path))
    doc.render(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated: {output_path} (template: {template_path.name})")


def generate_from_json(data: dict, output_path: Path) -> None:
    """Create a document from JSON data without a template."""
    if not _ensure_package("docx", "python-docx"):
        print("Error: python-docx is required.", file=sys.stderr)
        sys.exit(1)
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    if data.get("title"):
        doc.add_heading(data["title"], level=0)

    # Metadata fields
    for key in ("author", "date", "subtitle"):
        if data.get(key):
            doc.add_paragraph(f"{key.title()}: {data[key]}")

    # Summary
    if data.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data["summary"])

    # Sections
    for section in data.get("sections", []):
        if section.get("heading"):
            doc.add_heading(section["heading"], level=1)
        if section.get("body"):
            doc.add_paragraph(section["body"])

    # Items as table
    items = data.get("items", [])
    if items and isinstance(items[0], dict):
        doc.add_heading("Items", level=1)
        cols = list(items[0].keys())
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        for i, col_name in enumerate(cols):
            table.rows[0].cells[i].text = col_name.title()
        for item in items:
            row_cells = table.add_row().cells
            for i, col_name in enumerate(cols):
                row_cells[i].text = str(item.get(col_name, ""))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated: {output_path}")


def generate_from_markdown(md_text: str, output_path: Path) -> None:
    """Create a document from Markdown content."""
    if not _ensure_package("docx", "python-docx"):
        print("Error: python-docx is required.", file=sys.stderr)
        sys.exit(1)
    from docx import Document

    doc = Document()
    for line in md_text.split("\n"):
        stripped = line.rstrip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0 if not doc.paragraphs else 1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("1. ") or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == "."):
            doc.add_paragraph(stripped.split(". ", 1)[-1], style="List Number")
        elif stripped:
            doc.add_paragraph(stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Word document from JSON or Markdown input."
    )
    parser.add_argument(
        "input",
        type=Path,
        help="Path to JSON or Markdown input file."
    )
    parser.add_argument(
        "-t", "--template",
        type=Path,
        default=None,
        help="Path to a .docx template with Jinja2 tags. Used with JSON input only."
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=Path("output.docx"),
        help="Output file path (default: output.docx)."
    )
    args = parser.parse_args()

    if not args.input.is_file():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    if args.template and not args.template.is_file():
        print(f"Error: Template file not found: {args.template}", file=sys.stderr)
        sys.exit(1)

    if args.input.suffix.lower() == ".json":
        data = load_json(args.input)
        if args.template:
            generate_with_template(data, args.template, args.output)
        else:
            generate_from_json(data, args.output)
    else:
        if args.template:
            print("Warning: --template is ignored for Markdown input.", file=sys.stderr)
        md_text = args.input.read_text(encoding="utf-8")
        generate_from_markdown(md_text, args.output)


if __name__ == "__main__":
    main()
